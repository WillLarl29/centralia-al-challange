from pathlib import Path

from docx import Document as DocxDocument

from .base import BaseLoader, RawDocument


class DocxLoader(BaseLoader):
    def load(self, path: Path) -> list[RawDocument]:
        doc = DocxDocument(str(path))

        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        content = "\n".join(parts).strip()
        if not content:
            return []

        return [RawDocument(content=content, metadata={"source": path.name, "type": "docx"})]
